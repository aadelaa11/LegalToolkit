import streamlit as st
from openai import OpenAI
from datetime import datetime, timedelta
from fpdf import FPDF
from io import BytesIO
import json
import base64
from docx import Document
from PyPDF2 import PdfReader 
from streamlit_drawable_canvas import st_canvas

st.set_page_config(page_title="Integrated AI Legal Demand Letter Generator with Predictive Analysis", layout="wide")

st.markdown(
    """
    <style>
    div.stButton > button {
        background-color: white !important;
        color: black !important;
        border: 1px solid #ccc !important;
    }
    div.stButton > button:hover {
        background-color: #f0f0f0 !important;
        color: black !important;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# API Key Configuration
st.sidebar.title("API Key Configuration")
api_key = st.sidebar.text_input("Enter your OpenAI API Key", type="password", placeholder="sk-...")
if not api_key:
    st.sidebar.warning("Please enter your OpenAI API Key to proceed.")
    st.stop()
client = OpenAI(api_key=api_key)

st.markdown("<h1 style='text-align:center; color:#000; font-weight:700;'>Integrated Legal Demand Letter Generator with Predictive Analysis </h1>", unsafe_allow_html=True)
st.markdown("---")

# Risk & Contract Analysis Functions
def analyse_contract_formation(claim_reason, evidence_summary, evidence_strength, client):
    prompt = f"""
You are an Australian contract lawyer. Analyse the facts below to determine if a valid contract likely exists
and the probability of recovering the debt pre-litigation.

Assess each of the six essential elements:
1. Offer and acceptance
2. Intention to create legal relations 
3. Consideration
4. Legal Capacity
5. Consent
6. Legality

Facts / Claim reason:
{claim_reason}

Evidence summary:
{evidence_summary[:2000]}

User-rated evidence strength: {evidence_strength}/10
"""

    messages = [
        {"role": "system", "content": "You are a senior Australian contract law analyst."},
        {"role": "user", "content": prompt}
    ]
    response = client.chat.completions.create(model="gpt-5", messages=messages)
    return response.choices[0].message.content.strip()


def assess_risk_level(claim_amount, claim_reason, state, evidence_summary, evidence_strength, client):
    base_prompt = (
        f"Given the claim amount AUD {claim_amount:.2f}, reason: '{claim_reason}', "
        f"state: {state}, and evidence summary: '''{evidence_summary}''', "
        "assess the risk of successful debt recovery (scale 1–10) with a brief explanation."
    )
    base_messages = [
        {"role": "system", "content": "You are a legal analyst specialised in Australian debt recovery."},
        {"role": "user", "content": base_prompt}
    ]
    base_response = client.chat.completions.create(model="gpt-5", messages=base_messages)
    base_result = base_response.choices[0].message.content.strip()

    contract_analysis = analyse_contract_formation(claim_reason, evidence_summary, evidence_strength, client)

    return (
        "## Combined Legal Risk & Contract Formation Analysis\n\n"
        f"### Debt Recovery Risk Assessment\n{base_result}\n\n"
        f"{contract_analysis}"
    )

# Step 1: Demand Letter Details
st.header("Step 1: Demand Letter Details")

col1, col2 = st.columns([3, 2])
with col1:
    creditor_name = st.text_input("Your Name/Company (Creditor)", key="creditor_name")
    debtor_name = st.text_input("Debtor's Full Name or Company", key="debtor_name")
    claim_amount = st.number_input("Claim Amount (AUD)", min_value=0.0, step=0.01, format="%.2f", key="claim_amount")
    claim_reason = st.text_area("Reason for Claim", height=120, key="claim_reason")
with col2:
    response_days = st.number_input("Number of Days to Respond", min_value=1, max_value=60, value=14, key="response_days")
    letter_tone = st.selectbox("Letter Tone", ["Professional", "Firm", "Conciliatory", "Formal"], key="letter_tone")
    tone_intensity = st.slider("Tone Intensity", 1, 10, 5, key="tone_intensity")

st.markdown("---")

# Step 2: Upload Evidence (UPDATED SECTION)
st.header("Step 2: Upload Evidence")

uploaded_files = st.file_uploader(
    "Upload supporting evidence (TXT, PDF, DOCX)", 
    accept_multiple_files=True, 
    key="evidence_files"
)

def extract_text_from_files(files):
    """Extract readable text from TXT, DOCX, and PDF files."""
    extracted = ""
    for file in files:
        try:
            if file.type == "text/plain":
                extracted += file.getvalue().decode("utf-8") + "\n\n"
            elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                doc = Document(file)
                text = "\n".join([para.text for para in doc.paragraphs])
                extracted += text + "\n\n"
            elif file.type == "application/pdf":
                pdf_reader = PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() or ""
                extracted += text + "\n\n"
            else:
                extracted += f"[{file.name} - unsupported file type]\n\n"
        except Exception as e:
            extracted += f"[Error reading {file.name}: {str(e)}]\n\n"
    return extracted.strip()

if uploaded_files:
    extracted_text = extract_text_from_files(uploaded_files)
    st.session_state.extracted_text = extracted_text

    if extracted_text:
        st.markdown("**Evidence extracted successfully. Preview below:**")
        st.text_area("Extracted Evidence Text (first 1500 characters)", extracted_text[:1500], height=250)
    else:
        st.warning("⚠️ No readable text could be extracted. Please upload TXT, DOCX, or PDF files.")
else:
    extracted_text = ""
    st.session_state.extracted_text = ""
    st.info(" Upload one or more evidence files to strengthen your analysis.")

st.markdown("---")

#Step 3: Negotiation Options
st.header("Step 3: Negotiation Options")

offer_payment_plan = st.checkbox("Offer Payment Plan Option", key="payment_plan")
suggest_mediation = st.checkbox("Suggest Mediation", key="mediation")
demand_full_payment = st.checkbox("Demand Full Payment Upfront", value=True, key="full_payment")

#Step 4: Predictive Legal Analysis
st.markdown("---")
st.header("Step 4: Predictive Legal Analysis – Recovery, Contract Validity, and Litigation Outlook")

st.markdown("""
This step delivers an in-depth **Predictive Legal Analysis** powered by AI.  
It evaluates your matter based on the following faactors:

- Key strengths and weaknesses of your position  
- Whether a valid contract likely exists and how this may assist you in recovering your claim 
- Relevant Australian legal precedents  
- Probabilities of success pre- and post-litigation  
- Strategic recommendations and next steps  

""")

#Button to run analysis
run_button = st.button("Run Predictive Legal Analysis")

if run_button:
    with st.spinner("Analysing claim details, contract formation, and recovery potential..."):
        client = OpenAI(api_key=api_key)

        def run_predictive_analysis(claim_amount, claim_reason, state, evidence_summary, client):
            prompt = f"""
You are a senior Australian legal expert specialising in contract law, debt recovery, and pre-litigation advice.

Based on the following details, produce a **well-structured, concise, and clear predictive legal analysis** with professional formatting using clear headings, bullet points, and line spacing between sections.

---
FACTS:
- Claim Amount: AUD {claim_amount:.2f}
- Reason for Claim: {claim_reason}
- Jurisdiction: {state}
- Evidence Summary: {evidence_summary[:2000]}
---

Your output must include these clearly separated sections (with headings and blank lines between them):

## 1. Overview of the Matter
Briefly summarise the dispute and key issues.

## 2. Strengths and Weaknesses
List and explain:
- **Strengths:** (bullet points)
- **Weaknesses:** (bullet points)

## 3. Contract Formation Assessment
For each element, provide a separate bullet line formatted as:
**Element:** Yes/No/Unclear – short justification + relevant Australian case if applicable  
Include line breaks between each element.

Elements:
- Offer and acceptance 
- Intention to create legal relations   
- Consideration  
- Legal capacity
- Consent 
- Legality  

## 4. Predictive Risk Analysis
Show as:
- **Pre-litigation Recovery Probability:** xx%  
- **Litigation Success Probability:** xx%  
Provide concise justifications for both, with 2–3 explanatory sentences.

## 5. Expected Litigation Outlook
Summarise what to expect if litigation proceeds:
- Likely venue or tribunal  
- Timelines and procedural notes  
- Possible defences  
- Estimated costs  
- Settlement likelihood  

## 6. Relevant Legal Precedents
List 3–5 key Australian cases.  
Use bullet points formatted as:  
*Case Name (Year, Court)* – 1–2 line relevance summary.

## 7. Strategic Insights and Recommendations
Offer 3–5 actionable recommendations to improve position, manage risk, and prepare effectively.

---

Ensure clean line spacing and Markdown formatting for readability.
Avoid merging multiple points into a single paragraph.
"""

            messages = [
                {"role": "system", "content": "You are a senior Australian lawyer providing client-ready analysis."},
                {"role": "user", "content": prompt}
            ]
            response = client.chat.completions.create(model="gpt-5", messages=messages)
            return response.choices[0].message.content.strip()

        predictive_result = run_predictive_analysis(
            claim_amount=st.session_state.claim_amount,
            claim_reason=st.session_state.claim_reason,
            state=st.session_state.get("state", "Victoria"),
            evidence_summary=st.session_state.extracted_text,
            client=client
        )

        st.session_state["predictive_result"] = predictive_result

        st.success("Predictive Legal Analysis Complete")

        tabs = st.tabs(["Overview", "Strengths and Weaknesses", "Contract Formation", "Risk Analysis", "Litigation Outlook", "Legal Precedents", "Strategic Insights"])

        with tabs[0]:
            overview_content = predictive_result.split("## 2. Strengths and Weaknesses")[0]
            st.markdown(overview_content)

        with tabs[1]:
            st.markdown("### 2. Strengths and Weaknesses")
            strengths_weaknesses_content = predictive_result.split("## 3. Contract Formation Assessment")[0].split("## 2. Strengths and Weaknesses")[1]
            st.markdown(strengths_weaknesses_content)

        with tabs[2]:
            st.markdown("### 3. Contract Formation Assessment")
            contract_formation_content = predictive_result.split("## 4. Predictive Risk Analysis")[0].split("## 3. Contract Formation Assessment")[1]
            st.markdown(contract_formation_content)

        with tabs[3]:
            st.markdown("### 4. Predictive Risk Analysis")
            risk_analysis_content = predictive_result.split("## 5. Expected Litigation Outlook")[0].split("## 4. Predictive Risk Analysis")[1]
            st.markdown(risk_analysis_content)

        with tabs[4]:
            st.markdown("### 5. Expected Litigation Outlook")
            litigation_outlook_content = predictive_result.split("## 6. Relevant Legal Precedents")[0].split("## 5. Expected Litigation Outlook")[1]
            st.markdown(litigation_outlook_content)

        with tabs[5]:
            st.markdown("### 6. Relevant Legal Precedents")
            legal_precedents_content = predictive_result.split("## 7. Strategic Insights and Recommendations")[0].split("## 6. Relevant Legal Precedents")[1]
            st.markdown(legal_precedents_content)

        with tabs[6]:
            st.markdown("### 7. Strategic Insights and Recommendations")
            strategic_insights_content = predictive_result.split("## 7. Strategic Insights and Recommendations")[1]
            st.markdown(strategic_insights_content)

# Step 5: Generate Letter
st.markdown("---")
st.header("Generate Letter of Demand & Legal Analysis Summary")

include_interest = st.checkbox("Include Penalty Interest Clause", value=True, key="include_interest")
include_legal_costs = st.checkbox("Include Warning About Legal Costs", value=True, key="include_costs")

def get_legislation(state):
    laws = {
        "Victoria": "Victoria Civil and Administrative Tribunal (VCAT) Act 1998",
        "New South Wales": "Civil and Administrative Tribunal Act 2013",
        "Queensland": "Queensland Civil and Administrative Tribunal Act 2009",
    }
    return laws.get(state, "Relevant state legislation not found.")

def calculate_penalty_interest(claim_amount, overdue_days):
    rate = 0.08 / 365
    return claim_amount * rate * overdue_days

def generate_demand_letter(prompt, client):
    messages = [
        {"role": "system", "content": "You are a legal expert drafting formal Australian demand letters."},
        {"role": "user", "content": prompt}
    ]
    response = client.chat.completions.create(model="gpt-4o", messages=messages)
    return response.choices[0].message.content.strip()

def generate_letter_prompt(inputs):
    return f"""
Draft a {inputs['letter_tone']} demand letter under Australian {inputs['state']} law.
Creditor: {inputs['creditor_name']}
Debtor: {inputs['debtor_name']}
Claim amount: AUD {inputs['claim_amount']:.2f}
Reason: {inputs['claim_reason']}
Response deadline: {inputs['response_days']} days.
Include penalty interest: {inputs['include_interest']}
Include legal cost warning: {inputs['include_legal_costs']}
Evidence summary: {inputs['evidence_summary'][:1000]}
Negotiation options:
- Payment plan: {inputs['payment_plan']}
- Mediation: {inputs['mediation']}
- Demand full payment: {inputs['full_payment']}
Relevant legislation: {get_legislation(inputs['state'])}
"""

if st.button("Generate Letter of Demand Document"):
    with st.spinner("Preparing demand letter..."):
        today = datetime.today().date()
        overdue_days = 30
        penalty_interest = calculate_penalty_interest(claim_amount, overdue_days)
        inputs = {
            "creditor_name": creditor_name,
            "debtor_name": debtor_name,
            "claim_amount": claim_amount,
            "claim_reason": claim_reason,
            "response_days": response_days,
            "letter_tone": letter_tone,
            "state": "Victoria",
            "include_interest": include_interest,
            "include_legal_costs": include_legal_costs,
            "penalty_interest": penalty_interest,
            "payment_plan": offer_payment_plan,
            "mediation": suggest_mediation,
            "full_payment": demand_full_payment,
            "evidence_summary": extracted_text,
            "tone_intensity": tone_intensity
        }
        prompt = generate_letter_prompt(inputs)
        letter_text = generate_demand_letter(prompt, client)

        doc = Document()
        doc.add_heading('Letter of Demand', 0)
        doc.add_paragraph(letter_text)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.download_button(
            label="Download Demand Letter as Word Document",
            data=buffer,
            file_name="demand_letter.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

if st.button("Generate Legal Analysis Summary"):
    if 'predictive_result' not in st.session_state:
        st.warning("Please run the Predictive Legal Analysis first.")
    else:
        with st.spinner("Compiling your legal analysis summary..."):
            doc = Document()
            doc.add_heading('Legal Analysis Summary Report', 0)
            doc.add_heading('Claim Overview', level=1)
            doc.add_paragraph(f"Creditor: {creditor_name}")
            doc.add_paragraph(f"Debtor: {debtor_name}")
            doc.add_paragraph(f"Claim Amount: AUD {claim_amount:.2f}")
            doc.add_paragraph(f"Reason for Claim: {claim_reason}")
            doc.add_paragraph(f"Jurisdiction: Victoria")
            doc.add_paragraph(f"Response Period: {response_days} days")
            doc.add_heading('Predictive Legal Analysis', level=1)
            predictive_result = st.session_state.get("predictive_result", "")
            if predictive_result:
                doc.add_paragraph(predictive_result)
            else:
                doc.add_paragraph("No predictive analysis available yet.")
            doc.add_heading('Evidence Summary (User Upload)', level=1)
            if extracted_text:
                doc.add_paragraph(extracted_text[:1500])  # Include the first 1500 characters of evidence
            else:
                doc.add_paragraph("No evidence documents uploaded.")
            doc.add_heading('Strategic Insights and Recommendations', level=1)
            doc.add_paragraph(
                "Based on the predictive analysis, we suggest the following strategies and next steps:"
            )
            doc.add_paragraph(
                "- Assess the viability of settlement before initiating litigation."
            )
            doc.add_paragraph(
                "- Evaluate alternative dispute resolution options such as mediation."
            )
            doc.add_paragraph(
                "- Prepare to escalate the case to the relevant tribunal or court if recovery is unlikely without legal action."
            )
            doc.add_heading('Relevant Case References', level=1)
            if 'precedents' in st.session_state:
                precedents = st.session_state.get('precedents', [])
                for case in precedents:
                    doc.add_paragraph(f"- {case}")
            else:
                doc.add_paragraph("No case references generated yet.")

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            st.download_button(
                label="Download Legal Analysis Summary (Word)",
                data=buffer,
                file_name="legal_analysis_summary.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

st.markdown("---")
# Step 6: Interactive Legal Q&A
st.header("Legal Q&A (Interactive)")

# Ensure predictive results exist
if "predictive_result" not in st.session_state:
    st.info("Run the Predictive Legal Analysis first to provide context for Q&A.")
else:
    # Populate full_ai_results with predictive_result if not already done
    if "full_ai_results" not in st.session_state:
        st.session_state["full_ai_results"] = {"predictive_result": st.session_state["predictive_result"]}

    user_question = st.text_input("Enter your legal question about this case:")

    if st.button("Ask AI", key="qa_button"):
        if not user_question.strip():
            st.warning("Please type a question.")
        else:
            # Initialize Q&A history if not present
            if "qa_history" not in st.session_state:
                st.session_state["qa_history"] = []

            with st.spinner("Generating answer..."):
                # Combine all AI outputs in context
                context = "\n\n".join(st.session_state["full_ai_results"].values())

                qa_prompt = f"""
You are a senior Australian legal expert. Using the context below, answer the user's question concisely.
Context:
{context}

Question:
{user_question}
"""

                response = client.chat.completions.create(
                    model="gpt-5",
                    messages=[
                        {"role": "system", "content": "You are a senior Australian lawyer answering questions."},
                        {"role": "user", "content": qa_prompt}
                    ]
                )
                answer = response.choices[0].message.content.strip()

                # Save to history
                st.session_state["qa_history"].append({"question": user_question, "answer": answer})

    # Display Q&A history
    if "qa_history" in st.session_state and st.session_state["qa_history"]:
        st.markdown("### Response")
        for qa in reversed(st.session_state["qa_history"]):
            st.markdown(f"**Q:** {qa['question']}")
            st.markdown(f"**A:** {qa['answer']}")
            st.markdown("---")

# Step 7: Deadline Calculator
st.header("Additional Tools: Response Deadline Calculator")

base_date = st.date_input("Start Date", value=datetime.today())
deadline_days = st.number_input("Days to Respond", 1, 60, value=response_days, key="deadline_days")
include_weekends = st.checkbox("Count weekends?", True)

def calculate_deadline(start_date, days, include_weekends):
    date = start_date
    added = 0
    while added < days:
        date += timedelta(days=1)
        if include_weekends or date.weekday() < 5:
            added += 1
    return date

deadline = calculate_deadline(base_date, deadline_days, include_weekends)
st.success(f"Calculated response deadline: **{deadline.strftime('%A, %d %B %Y')}**")